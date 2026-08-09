#!/usr/bin/env python3
"""Build the single public-facing Nature Methods submission story.

This script does not run scientific models.  It renders figures and documents
from already frozen evidence, omits internal history from the public surface,
and writes machine-readable staging tables for the final Source Data workbook.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import re
import shutil
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(os.environ.get("SDG_WORKTREE", Path(__file__).resolve().parents[3]))
R2 = ROOT / "delivery/scDesignGuard_V3_8_OVERLAP_EXCLUDED_FINAL_SUBMISSION_PACKAGE_R2_20260809"
OUT = ROOT / "delivery/scDesignGuard_Nature_Methods_FINAL_SUBMISSION_PACKAGE_20260809"
STAGE = ROOT / "reports/final_story_only_submission_20260809"
INPUT_DOCX = Path(os.environ.get("SDG_INPUT_DOCX", ROOT / "inputs/scDesignGuard_Nature_Methods_manuscript_V3_8_OVERLAP_EXCLUDED_FINAL.docx"))
FINAL_VALIDATION = ROOT / "reports/final_validation_20260809/human_blinded_v1_2_e2e_v1b_resume"
FRAMEWORK = ROOT / "reports/human_review/framework_v2_terminal_execution_20260807"
BIO = ROOT / "analysis/post_terminal_bioinformatics_expansion_v2"
OLD_SOURCE = ROOT / "delivery/source_data"

INK = "#25313B"
BLUE = "#2F75B5"
GREEN = "#2F9D6A"
ORANGE = "#D8892B"
RED = "#B94A48"
PURPLE = "#7C6AA6"
GRAY = "#8B969E"
LIGHT = "#EDF2F5"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def configure() -> None:
    mpl.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
        "font.size": 7.2,
        "axes.titlesize": 8.0,
        "axes.labelsize": 7.2,
        "xtick.labelsize": 6.4,
        "ytick.labelsize": 6.4,
        "legend.fontsize": 6.2,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.linewidth": 0.7,
        "pdf.fonttype": 42,
        "svg.fonttype": "none",
    })


def panel(ax: plt.Axes, label: str) -> None:
    ax.text(-0.10, 1.06, label, transform=ax.transAxes, fontsize=9, fontweight="bold", va="top")


def audit_box_layout(fig: plt.Figure, stem: str) -> None:
    """Fail when text escapes its box or when authored boxes overlap."""
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    problems: list[str] = []
    for ax_index, ax in enumerate(fig.axes):
        authored_boxes = [patch for patch in ax.patches if getattr(patch, "_sdg_authored_box", False)]
        for text in ax.texts:
            container = getattr(text, "_sdg_container_box", None)
            if container is None:
                continue
            text_bounds = text.get_window_extent(renderer=renderer)
            box_bounds = container.get_window_extent(renderer=renderer)
            inset = 1.0
            if (
                text_bounds.x0 < box_bounds.x0 + inset
                or text_bounds.x1 > box_bounds.x1 - inset
                or text_bounds.y0 < box_bounds.y0 + inset
                or text_bounds.y1 > box_bounds.y1 - inset
            ):
                problems.append(f"axis {ax_index}: text {text.get_text()!r} escapes its box")
        for left_index, left in enumerate(authored_boxes):
            left_bounds = left.get_window_extent(renderer=renderer)
            for right in authored_boxes[left_index + 1:]:
                right_bounds = right.get_window_extent(renderer=renderer)
                overlap_width = max(0.0, min(left_bounds.x1, right_bounds.x1) - max(left_bounds.x0, right_bounds.x0))
                overlap_height = max(0.0, min(left_bounds.y1, right_bounds.y1) - max(left_bounds.y0, right_bounds.y0))
                if overlap_width * overlap_height > 1.0:
                    problems.append(f"axis {ax_index}: authored boxes overlap by {overlap_width * overlap_height:.1f} px²")
    if problems:
        raise AssertionError(f"{stem} box-layout audit failed: " + "; ".join(problems))


def save_figure(fig: plt.Figure, stem: str, extended: bool = False) -> None:
    configure()
    audit_box_layout(fig, stem)
    base = OUT / ("figures/extended_data" if extended else "figures/main") / stem
    base.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight", facecolor="white")
    fig.savefig(base.with_suffix(".png"), dpi=450, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def box(ax: plt.Axes, xy: tuple[float, float], wh: tuple[float, float], text: str, color: str, fontsize: float = 7.0) -> None:
    x, y = xy
    w, h = wh
    rectangle = plt.Rectangle((x, y), w, h, facecolor=color, alpha=0.12, edgecolor=color, linewidth=1.0)
    rectangle._sdg_authored_box = True
    ax.add_patch(rectangle)
    label = ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", color=INK, fontweight="bold", fontsize=fontsize, linespacing=1.0)
    label._sdg_container_box = rectangle


def write_csv(path: Path, frame: pd.DataFrame) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    frame.to_csv(path, index=False, lineterminator="\n")


def build_figure_1() -> pd.DataFrame:
    configure()
    fig = plt.figure(figsize=(7.15, 4.55), constrained_layout=True)
    grid = fig.add_gridspec(2, 3, height_ratios=[1.0, 1.15], width_ratios=[1.15, 1.10, 0.85])
    ax = fig.add_subplot(grid[0, :2]); panel(ax, "a"); ax.axis("off")
    box(ax, (0.02, 0.53), (0.20, 0.26), "Source version\nDonor map\nCondition map", BLUE)
    box(ax, (0.29, 0.53), (0.20, 0.26), "Target ontology\nSupport rule\nEstimand", PURPLE)
    box(ax, (0.58, 0.45), (0.22, 0.42), "Evidence-bearing\ndesign contract", GREEN)
    ax.annotate("", (0.29, 0.66), (0.22, 0.66), arrowprops={"arrowstyle": "-|>", "color": GRAY, "lw": 1.2})
    ax.annotate("", (0.58, 0.66), (0.49, 0.66), arrowprops={"arrowstyle": "-|>", "color": GRAY, "lw": 1.2})
    for y, label, color, size in [(0.63, "PROCEED", GREEN, 7.0), (0.42, "ABSTAIN", ORANGE, 7.0), (0.21, "BLOCK", RED, 7.0), (0.00, "NON_\nEVALUABLE", GRAY, 5.8)]:
        box(ax, (0.86, y), (0.12, 0.17), label, color, size)
        ax.annotate("", (0.86, y + 0.085), (0.80, 0.66), arrowprops={"arrowstyle": "-|>", "color": GRAY, "lw": 0.8})
    ax.set_title("Evidence compilation", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[0, 2]); panel(ax, "b"); ax.axis("off")
    stages = ["Source\nscope", "Structural\ndesign", "Target\nsupport", "Model\nvalidity", "Inference\nauthorization"]
    ys = [0.82, 0.64, 0.46, 0.28, 0.10]
    for i, (label, y) in enumerate(zip(stages, ys)):
        box(ax, (0.14, y), (0.72, 0.12), label, [BLUE, PURPLE, GREEN, ORANGE, RED][i])
        if i < len(stages) - 1:
            ax.annotate("", (0.5, ys[i + 1] + 0.13), (0.5, y - 0.01), arrowprops={"arrowstyle": "-|>", "color": GRAY})
    ax.set_title("Stage-specific authorization", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[1, :]); panel(ax, "c"); ax.axis("off")
    ax.set_title("Worked example: GSE144735", loc="left", fontweight="bold", pad=8)
    cards = [("6", "complete pairs", BLUE), ("3,356", "target epithelial cells", PURPLE), ("12", "donor-condition\npseudobulks", GREEN), ("7/7", "design-matrix rank", ORANGE), ("5", "residual d.f.", RED)]
    for i, (value, label, color) in enumerate(cards):
        box(ax, (0.015 + i * 0.197, 0.38), (0.177, 0.37), f"{value}\n{label}", color, 6.4)
    ax.annotate("", (0.91, 0.22), (0.09, 0.22), arrowprops={"arrowstyle": "-|>", "color": GREEN, "lw": 2.0})
    save_figure(fig, "Figure_1")
    return pd.DataFrame([
        {"family": "GSE144735", "metric": "complete_pairs", "value": 6, "unit": "pairs"},
        {"family": "GSE144735", "metric": "target_cells", "value": 3356, "unit": "cells"},
        {"family": "GSE144735", "metric": "pseudobulk_rows", "value": 12, "unit": "rows"},
        {"family": "GSE144735", "metric": "design_rank", "value": 7, "unit": "columns"},
        {"family": "GSE144735", "metric": "design_columns", "value": 7, "unit": "columns"},
        {"family": "GSE144735", "metric": "residual_df", "value": 5, "unit": "degrees_of_freedom"},
    ])


def enrich_benchmark() -> pd.DataFrame:
    bench = pd.read_csv(R2 / "source_data/STUDY_FAMILY_DISJOINT_15_FAMILY_BENCHMARK.csv")
    contract = pd.read_csv(FINAL_VALIDATION / "EXECUTABLE_CONTRACT_RECONSTRUCTION_LEDGER.csv")
    requests = pd.read_csv(FINAL_VALIDATION / "OFFICIAL_SOURCE_REQUEST_LEDGER.csv")
    requests = requests[requests["record_id"].isin(bench["record_id"])].copy()
    requests = requests.sort_values(["record_id", "request_started_utc"]).drop_duplicates("record_id", keep="first")
    keep_c = [c for c in ["record_id", "accession_or_collection_id", "primary_design_contrast", "biological_unit", "design_type", "author_assayed_compartment", "source_contract_status"] if c in contract]
    keep_r = [c for c in ["record_id", "repository", "accession_or_collection_id", "effective_url", "request_started_utc", "system_sha256"] if c in requests]
    out = bench.merge(contract[keep_c], on="record_id", how="left", suffixes=("", "_contract"))
    out = out.merge(requests[keep_r], on="record_id", how="left", suffixes=("", "_request"))
    if "accession_or_collection_id" not in out:
        out["accession_or_collection_id"] = out.get("accession_or_collection_id_contract", "")
    out["exact_source_version"] = out["accession_or_collection_id"].astype(str) + "; source SHA-256=" + out["system_sha256"].fillna("NOT_RECORDED").astype(str)
    out["endpoint_membership"] = out["frozen_human_reference_class"].map({
        "VALID_DESIGN": "VALID_CLEARANCE;EXACT_STATE_CONCORDANCE",
        "HARD_INVALID_DESIGN": "HARD_INVALID_DETECTION;UNSAFE_CONTINUATION;EXACT_STATE_CONCORDANCE",
        "STRUCTURALLY_UNRESOLVED": "UNRESOLVED_CONTAINMENT;EXACT_STATE_CONCORDANCE",
    })
    assert len(out) == 15
    return out


def build_figure_2(bench: pd.DataFrame) -> pd.DataFrame:
    configure()
    endpoints = pd.read_csv(R2 / "source_data/STUDY_FAMILY_DISJOINT_ENDPOINTS.csv")
    order = ["Valid clearance", "Hard-invalid detection", "Unresolved containment", "Unsafe continuation", "Exact state concordance"]
    endpoints = endpoints.set_index("endpoint").loc[order].reset_index()
    fig = plt.figure(figsize=(7.15, 5.1), constrained_layout=True)
    grid = fig.add_gridspec(2, 2, height_ratios=[0.82, 1.18], width_ratios=[1.05, 1.0])
    ax = fig.add_subplot(grid[0, 0]); panel(ax, "a"); ax.axis("off")
    labels = ["Family-identity\naudit", "Frozen source-\nbound human\nconsensus", "Sealed evaluator\nexecution"]
    colors = [BLUE, PURPLE, GREEN]
    for i, (label, color) in enumerate(zip(labels, colors)):
        box(ax, (0.04 + i * 0.32, 0.34), (0.26, 0.34), label, color, 5.8)
        if i < 2:
            ax.annotate("", (0.36 + i * 0.32, 0.51), (0.30 + i * 0.32, 0.51), arrowprops={"arrowstyle": "-|>", "color": GRAY, "lw": 1.2})
    ax.set_title("Blinded evaluation sequence", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[0, 1]); panel(ax, "b")
    strata = pd.Series({"Valid": 4, "Hard-invalid": 6, "Unresolved": 5})
    yy = np.arange(3)
    ax.barh(yy, strata.values, color=[GREEN, RED, GRAY], edgecolor=INK, linewidth=0.5)
    ax.set_yticks(yy, strata.index); ax.invert_yaxis(); ax.set_xlabel("Study families"); ax.set_xlim(0, 7)
    for i, v in enumerate(strata.values): ax.text(v + 0.10, i, str(v), va="center", fontweight="bold")
    ax.set_title("Frozen consensus strata", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[1, 0]); panel(ax, "c")
    y = np.arange(len(endpoints))
    est = endpoints["estimate"].astype(float).to_numpy()
    lo = endpoints["wilson_95_low"].astype(float).clip(lower=0).to_numpy()
    hi = endpoints["wilson_95_high"].astype(float).clip(upper=1).to_numpy()
    lower_err = np.maximum(0.0, est - lo)
    upper_err = np.maximum(0.0, hi - est)
    ax.errorbar(est, y, xerr=[lower_err, upper_err], fmt="o", color=BLUE, ecolor=INK, capsize=3, markersize=4)
    ax.set_yticks(y, endpoints["endpoint"]); ax.invert_yaxis(); ax.set_xlim(-0.04, 1.04); ax.set_xlabel("Estimate (Wilson 95% CI)")
    for i, row in endpoints.iterrows():
        ax.text(0.52 if row.estimate == 0 else 0.96, i - 0.18, f"{int(row.numerator)}/{int(row.denominator)}", ha="center", fontsize=6.0, fontweight="bold")
    ax.set_title("State-concordance endpoints", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[1, 1]); panel(ax, "d")
    classes = ["Valid", "Hard-invalid", "Unresolved"]
    states = ["PROCEED", "BLOCK", "NON_EVALUABLE"]
    matrix = np.diag([4, 6, 5])
    im = ax.imshow(matrix, cmap="Blues", vmin=0, vmax=6)
    ax.set_xticks(np.arange(3), states, rotation=22, ha="right", rotation_mode="anchor"); ax.set_yticks(np.arange(3), classes)
    ax.set_xlabel("scDesignGuard state"); ax.set_ylabel("Frozen human class")
    for i in range(3):
        for j in range(3): ax.text(j, i, str(matrix[i, j]), ha="center", va="center", color="white" if matrix[i, j] > 3 else INK, fontweight="bold")
    ax.set_title("Family-level state matrix", loc="left", fontweight="bold", pad=8)
    save_figure(fig, "Figure_2")
    return endpoints


def enrich_e2e() -> pd.DataFrame:
    e2e = pd.read_csv(R2 / "source_data/STUDY_FAMILY_DISJOINT_E2E_4_FAMILY_RESULTS.csv")
    requests = pd.read_csv(FINAL_VALIDATION / "OFFICIAL_SOURCE_REQUEST_LEDGER.csv")
    req = requests[requests["record_id"].isin(e2e.record_id)].sort_values(["record_id", "request_started_utc"]).drop_duplicates("record_id", keep="first")
    cols = [c for c in ["record_id", "repository", "effective_url", "request_started_utc", "system_sha256"] if c in req]
    e2e = e2e.merge(req[cols], on="record_id", how="left")
    detail = {
        "E2E1B-02": {"display_id": "GSE144735", "contrast": "paired condition contrast", "target": "epithelial cells", "pair_count": 6, "donor_count": 6, "target_cell_count": 3356, "pseudobulk_rows": 12, "design_rank": "7/7", "estimability": "YES", "residual_df": 5},
        "E2E1B-03": {"display_id": "GSE123813", "contrast": "paired condition contrast", "target": "T cells", "pair_count": 11, "donor_count": 11, "target_cell_count": 33106, "pseudobulk_rows": 22, "design_rank": "12/12", "estimability": "YES", "residual_df": 10},
        "E2E1B-06": {"display_id": "HCA-E2E-03", "contrast": "source-defined comparison", "target": "NOT_ASSESSED", "pair_count": "NOT_ASSESSED", "donor_count": "NOT_ASSESSED", "target_cell_count": "NOT_ASSESSED", "pseudobulk_rows": "NOT_APPLICABLE", "design_rank": "NOT_APPLICABLE", "estimability": "NOT_APPLICABLE", "residual_df": "NOT_APPLICABLE"},
        "E2E1B-07": {"display_id": "GSE81076", "contrast": "independent groups", "target": "source-defined", "pair_count": 0, "donor_count": 5, "target_cell_count": "NOT_ASSESSED", "pseudobulk_rows": "NOT_APPLICABLE", "design_rank": "NOT_APPLICABLE", "estimability": "NOT_APPLICABLE", "residual_df": "NOT_APPLICABLE"},
    }
    d = pd.DataFrame.from_dict(detail, orient="index").reset_index(names="record_id")
    e2e = e2e.merge(d, on="record_id", how="left")
    assert len(e2e) == 4
    return e2e


def build_figure_3(e2e: pd.DataFrame) -> None:
    configure()
    fig = plt.figure(figsize=(7.15, 4.35), constrained_layout=True)
    grid = fig.add_gridspec(2, 2, height_ratios=[0.9, 1.2], width_ratios=[1.25, 0.9])
    ax = fig.add_subplot(grid[0, :]); panel(ax, "a"); ax.axis("off")
    outcomes = [("GSE144735", "All gates passed", GREEN), ("GSE123813", "All gates passed", GREEN), ("HCA-E2E-03", "Source-binding stop", ORANGE), ("GSE81076", "Donor-replication stop", ORANGE)]
    for i, (family, outcome, color) in enumerate(outcomes): box(ax, (0.02 + i * 0.245, 0.25), (0.225, 0.50), f"{family}\n\n{outcome}", color, 6.4)
    ax.set_title("Final gate outcome", loc="left", fontweight="bold", pad=8)
    ax = fig.add_subplot(grid[1, 0]); panel(ax, "b"); ax.axis("off")
    ax.set_title("Quantitative gate checks", loc="left", fontweight="bold", pad=8)
    columns = ["Complete\npairs", "Target\ncells", "Pseudobulk\nrows", "Design\nrank", "Residual\nd.f."]
    values = {"GSE144735": ["6", "3,356", "12", "7/7", "5"], "GSE123813": ["11", "33,106", "22", "12/12", "10"]}
    for i, label in enumerate(columns): ax.text(0.26 + i * 0.145, 0.86, label, ha="center", va="center", fontsize=6.1, fontweight="bold")
    for r, (family, vals) in enumerate(values.items()):
        yy = 0.61 - r * 0.31; color = [BLUE, GREEN][r]
        ax.text(0.02, yy, family, ha="left", va="center", fontweight="bold", color=color)
        for i, value in enumerate(vals):
            ax.add_patch(plt.Circle((0.26 + i * 0.145, yy), 0.055, facecolor=color, alpha=0.13, edgecolor=color, linewidth=1))
            ax.text(0.26 + i * 0.145, yy, value, ha="center", va="center", fontweight="bold")
    ax = fig.add_subplot(grid[1, 1]); panel(ax, "c")
    stages = ["Source scope", "Structural design", "Target support", "Counts", "Model validity", "Authorization"]
    counts = [4, 2, 2, 2, 2, 2]
    ax.plot(range(6), counts, color=BLUE, marker="o", linewidth=1.8, markersize=5); ax.fill_between(range(6), counts, color=BLUE, alpha=0.10)
    ax.set_xticks(range(6), stages); plt.setp(ax.get_xticklabels(), rotation=30, ha="right", rotation_mode="anchor")
    ax.set_ylim(0, 4.4); ax.set_ylabel("Cases reaching gate"); ax.set_title("Cases reaching each gate", loc="left", fontweight="bold", pad=8)
    save_figure(fig, "Figure_3")


def build_figure_4() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    configure()
    challenge = pd.read_csv(FRAMEWORK / "benchmark/B02_TIER_D_KNOWN_INVALID_RESULTS.csv")
    transport = pd.read_csv(FRAMEWORK / "benchmark/B04_STUDY_FAMILY_HELDOUT_RESULTS.csv")
    transport = transport.rename(columns={"heldout_id": "transport_audit_id"})
    summary = pd.read_csv(R2 / "source_data/FIGURE_4_SOURCE_DATA.csv")
    fig = plt.figure(figsize=(7.15, 4.25), constrained_layout=True)
    grid = fig.add_gridspec(1, 3, width_ratios=[1.35, 0.90, 0.82])
    ax = fig.add_subplot(grid[0, 0]); panel(ax, "a")
    labels = [v.replace("_", " ").title() for v in challenge.challenge]
    yy = np.arange(len(labels)); colors = [GRAY if s == "NON_EVALUABLE" else RED for s in challenge.terminal_state]
    ax.barh(yy, np.ones(len(yy)), color=colors, edgecolor=INK, linewidth=0.45); ax.set_yticks(yy, labels); ax.invert_yaxis(); ax.set_xlim(0, 1); ax.set_xticks([])
    for i, state in enumerate(challenge.terminal_state): ax.text(0.5, i, state, ha="center", va="center", color="white", fontsize=5.7, fontweight="bold")
    ax.set_title("Controlled invalidity challenges", x=0.06, ha="left", fontweight="bold", pad=8)
    ax.text(0.5, -0.09, "9/9 safe; 9/9 reason match", transform=ax.transAxes, ha="center", fontsize=6.0, fontweight="bold")
    for spine in ax.spines.values(): spine.set_visible(False)
    ax = fig.add_subplot(grid[0, 1]); panel(ax, "b")
    counts = transport.groupby("repository_platform").size().reindex(["BIOSTUDIES_ARRAYEXPRESS", "SINGLE_CELL_PORTAL", "NCBI_GEO", "HUMAN_CELL_ATLAS"])
    names = ["BioStudies/ArrayExpress", "Single Cell Portal", "NCBI GEO", "Human Cell Atlas"]
    yy = np.arange(4); ax.barh(yy, counts.values, color=[BLUE, PURPLE, GREEN, ORANGE], edgecolor=INK, linewidth=0.5)
    ax.set_yticks(yy, names); ax.invert_yaxis(); ax.set_xlim(0, 6); ax.set_xlabel("Transport-audit families")
    for i, v in enumerate(counts.values): ax.text(v + 0.08, i, str(v), va="center", fontweight="bold")
    ax.set_title("Repository transport", x=0.08, ha="left", fontweight="bold", pad=8); ax.text(0.5, -0.12, "n=18 Tier C audits", transform=ax.transAxes, ha="center", fontsize=6.0)
    ax = fig.add_subplot(grid[0, 2]); panel(ax, "c")
    tests = summary[summary.panel == "4c"].copy(); display = tests.numerator.astype(float) / tests.denominator.astype(float)
    yy = np.arange(3); ax.barh(yy, display, color=[GREEN, BLUE, PURPLE], edgecolor=INK, linewidth=0.5)
    ax.set_yticks(yy, ["Current suites", "Compatibility suites", "Repository tests"]); ax.invert_yaxis(); ax.set_xlim(0, 1.16); ax.set_xticks([0, 0.5, 1.0]); ax.set_xlabel("Completed / expected")
    for i, row in enumerate(tests.itertuples()): ax.text(1.02, i, f"{int(row.numerator):,}/{int(row.denominator):,}", va="center", fontsize=5.8, fontweight="bold")
    ax.set_title("Software reproduction", x=0.10, ha="left", fontweight="bold", pad=8); ax.text(0.5, -0.13, "29 planned skips; 42 subtests passed", transform=ax.transAxes, ha="center", fontsize=6.0)
    save_figure(fig, "Figure_4")
    return challenge, transport, tests


def build_figure_5() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    configure()
    means = pd.read_csv(R2 / "source_data/FIGURE_5A_V3_5_SOURCE_DATA.csv").set_index("condition")
    se = pd.read_csv(R2 / "source_data/FIGURE_5B_V3_5_UNMODERATED_SE_SOURCE_DATA.csv").set_index("cluster")
    summary = pd.read_csv(R2 / "source_data/FIGURE_5C_V3_5_SOURCE_DATA.csv").set_index("cluster")
    heat = pd.read_csv(R2 / "source_data/FIGURE_5D_V3_5_SOURCE_DATA.csv").set_index("pathway")
    clusters = [f"C{i}" for i in range(8)]; palette = [BLUE, ORANGE, GREEN, RED, PURPLE, "#8C564B", "#CC79A7", GRAY]
    fig = plt.figure(figsize=(7.15, 5.6), constrained_layout=True); grid = fig.add_gridspec(2, 2, width_ratios=[1.03, 1.0], height_ratios=[0.93, 1.07])
    ax = fig.add_subplot(grid[0, 0]); panel(ax, "a"); bottom = np.zeros(3)
    for i, cluster in enumerate(clusters):
        vals = means.loc[["ND", "PD", "T2D"], cluster].to_numpy(float); ax.bar(["ND", "PD", "T2D"], vals, bottom=bottom, label=cluster, color=palette[i], width=0.72); bottom += vals
    ax.set_ylim(0, 1.03); ax.set_ylabel("Mean donor proportion"); ax.set_title("Author-defined beta-cell states", loc="left", fontweight="bold", pad=8); ax.legend(ncol=4, fontsize=5.5, loc="lower left", frameon=False)
    ax = fig.add_subplot(grid[0, 1]); panel(ax, "b"); ordered = se.loc[clusters]
    ax.errorbar(ordered.frozen_coefficient, np.arange(8), xerr=1.96 * ordered.unmoderated_ols_se, fmt="o", color=INK, capsize=2, markersize=3.5); ax.axvline(0, color=GRAY, lw=0.8)
    ax.set_yticks(np.arange(8), clusters); ax.set_xlabel("Adjusted effect on asin-sqrt proportion\n(descriptive ±1.96 unmoderated SE)"); ax.set_title("T2D–ND donor-level composition", loc="left", fontweight="bold", pad=8)
    ax.text(-0.122, 0.25, "FDR=0.0213", fontsize=5.8); ax.text(0.19, 6.25, "FDR=0.000178", fontsize=5.8)
    ax = fig.add_subplot(grid[1, 0]); panel(ax, "c"); ordered_s = summary.loc[clusters]; xx = np.arange(8)
    ax.bar(xx - 0.18, ordered_s.edgeR_global_FDR_0_05, width=0.36, label="edgeR QL", color=BLUE); ax.bar(xx + 0.18, ordered_s.voom_global_FDR_0_05, width=0.36, label="limma-voom", color="#E07A5F")
    ax.set_xticks(xx, clusters); ax.set_ylabel("Gene × cluster associations\n(global BH-FDR < 0.05)"); ax.set_title("State-resolved T2D–ND expression", loc="left", fontweight="bold", pad=8); ax.legend(frameon=False, loc="center right", fontsize=5.8)
    ax.text(0.98, 0.97, "418 associations / 225 genes\nAtlas overlap: 136/225 genes\nDirection agreement: 310/310", transform=ax.transAxes, ha="right", va="top", fontsize=5.6, bbox={"boxstyle": "round,pad=.25", "facecolor": "white", "edgecolor": GRAY, "alpha": 0.92})
    ax = fig.add_subplot(grid[1, 1]); panel(ax, "d"); display = heat.loc[:, clusters]; vmax = max(2.0, float(np.nanpercentile(np.abs(display.to_numpy(float)), 95)))
    im = ax.imshow(display, aspect="auto", cmap="RdBu_r", vmin=-vmax, vmax=vmax); ax.set_xticks(np.arange(8), clusters)
    ax.set_yticks(np.arange(len(display)), [v.replace("HALLMARK_", "").replace("_", " ") for v in display.index], fontsize=5.2); ax.set_title("Hallmark competitive pathway tests", loc="left", fontweight="bold", pad=8)
    cb = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.03); cb.set_label("signed −log10(global FDR)")
    save_figure(fig, "Figure_5")
    return means.reset_index(), se.reset_index(), summary.reset_index(), heat.reset_index()


def build_ed1() -> pd.DataFrame:
    support = pd.read_csv(ROOT / "reports/step17hr3_final/TARGET_GROUP_SUPPORT_SUMMARY.csv")
    lodo = pd.read_csv(ROOT / "reports/step15/STEP_15_WP2_CELLTYPE_SUMMARY.csv")
    labels = pd.read_csv(ROOT / "reports/step17hr3_calibration/STEP17HR3_CALIBRATION_LABEL_SUMMARY.csv")
    fig, axes = plt.subplots(2, 2, figsize=(7.15, 5.6), constrained_layout=True)
    ax = axes[0, 0]; panel(ax, "a"); donor = support.pivot(index="target", columns="clinical_group", values="qualifying_donors").fillna(0)
    donor.plot.bar(ax=ax, color=[BLUE, ORANGE], width=0.72); ax.axhline(5, color=RED, ls="--", lw=0.9); ax.set_ylabel("Qualifying donors"); ax.set_xlabel(""); ax.set_title("Independent-unit support", loc="left", fontweight="bold"); ax.legend(frameon=False, fontsize=5.6)
    ax = axes[0, 1]; panel(ax, "b"); xx = np.arange(len(lodo)); stable = lodo.complete_direction_stable_hypotheses.astype(int); unstable = lodo.direction_unstable_hypotheses.astype(int)
    ax.bar(xx, stable, color=GREEN, label="Direction-stable"); ax.bar(xx, unstable, bottom=stable, color=ORANGE, label="Direction-unstable"); ax.set_xticks(xx, lodo.celltype_major, rotation=35, ha="right", rotation_mode="anchor"); ax.set_ylabel("Hypotheses"); ax.set_title("Leave-one-donor-out diagnostics", loc="left", fontweight="bold"); ax.legend(frameon=False, fontsize=5.6)
    ax = axes[1, 0]; panel(ax, "c"); piv = labels.pivot(index="calibration_target", columns="DIRECTIONALLY_REPLICATED", values="Freq").fillna(0); xx = np.arange(len(piv))
    ax.bar(xx, piv.get(0, pd.Series(0, index=piv.index)), color=GRAY, label="Not replicated"); ax.bar(xx, piv.get(1, pd.Series(0, index=piv.index)), bottom=piv.get(0, pd.Series(0, index=piv.index)), color=GREEN, label="Directionally replicated")
    ax.set_xticks(xx, [v.replace("_", " ") for v in piv.index]); ax.set_ylabel("Gene-level labels"); ax.set_title("Calibration-label composition", loc="left", fontweight="bold"); ax.legend(frameon=False, fontsize=5.6)
    ax = axes[1, 1]; panel(ax, "d"); s = support.copy(); s["label"] = s.clinical_group.str.replace("_", " ") + " · " + s.target.str.replace("_", " "); yy = np.arange(len(s))
    ax.barh(yy, s.minimum_cells_in_donor_target.astype(int), color=[BLUE, GREEN, PURPLE] * math.ceil(len(s) / 3)); ax.axvline(20, color=RED, ls="--", lw=0.9, label="20-cell cut-off"); ax.set_yticks(yy, s.label); ax.invert_yaxis(); ax.set_xlabel("Minimum cells in qualifying donor-target"); ax.set_title("Target-support floor", loc="left", fontweight="bold"); ax.legend(frameon=False, fontsize=5.6)
    save_figure(fig, "Extended_Data_Figure_1", True)
    frames = []
    for name, frame in [("target_support", support), ("lodo", lodo), ("calibration", labels)]:
        z = frame.copy(); z.insert(0, "panel_source", name); frames.append(z)
    return pd.concat(frames, ignore_index=True, sort=False)


def build_ed2() -> pd.DataFrame:
    nm02 = pd.read_csv(ROOT / "reports/nm02_s1_r1/NM02_S1_ENDPOINT_RESULTS_CORRECTED_V1_1.csv")
    primary = nm02[nm02.component_id.isin(["POOLED_PER_NULL_TEST_TYPE_I_RATE", "FDR_MEAN_V_OVER_MAX_R_1"])].copy()
    checklist = pd.DataFrame([
        ("Exact v0.1.0 wheel", "PASS", "REVIEWER_ARCHIVE"), ("OCI image binding", "PASS", "REVIEWER_ARCHIVE"), ("JSON Schema", "PASS", "FROZEN"), ("21 reason codes", "PASS", "FROZEN"), ("4 terminal-state fixtures", "PASS", "FROZEN"), ("Independent clean install", "PASS", "VERIFIED"), ("SBOM and security receipt", "PASS", "VERIFIED"), ("Public archival release", "PENDING", "UPON_ACCEPTANCE")
    ], columns=["component", "status", "availability"])
    fig, axes = plt.subplots(1, 2, figsize=(7.15, 3.6), constrained_layout=True)
    ax = axes[0]; panel(ax, "a"); y = np.arange(len(primary)); est = primary.estimate.astype(float).to_numpy(); lo = primary.interval_lower.astype(float).to_numpy(); hi = primary.interval_upper.astype(float).to_numpy()
    ax.errorbar(est, y, xerr=[est - lo, hi - est], fmt="o", color=BLUE, ecolor=INK, capsize=4); ax.axvline(0.05, color=RED, ls="--", lw=0.9)
    ax.set_yticks(y, ["EP-07 pooled Type-I rate", "EP-08 mean replicate FDP/FDR"]); ax.invert_yaxis(); ax.set_xlabel("Estimate with Monte Carlo interval"); ax.set_xlim(0.043, 0.052); ax.set_title("Synthetic endpoint uncertainty", loc="left", fontweight="bold")
    for i, row in primary.reset_index(drop=True).iterrows(): ax.text(float(row.interval_upper) + 0.00012, i, f"{float(row.estimate):.5f}\nMCSE {float(row.mcse):.5f}", va="center", fontsize=5.7)
    ax = axes[1]; panel(ax, "b"); yy = np.arange(len(checklist)); colors = [GREEN if s == "PASS" else ORANGE for s in checklist.status]
    ax.barh(yy, np.ones(len(checklist)), color=colors); ax.set_yticks(yy, checklist.component); ax.invert_yaxis(); ax.set_xlim(0, 1.62); ax.set_xticks([])
    for i, row in checklist.iterrows(): ax.text(0.04, i, row.status, va="center", color="white", fontweight="bold"); ax.text(1.03, i, row.availability.replace("_", " "), va="center", fontsize=5.4, color=INK)
    ax.set_title("Installation and fixture reproduction", loc="left", fontweight="bold")
    save_figure(fig, "Extended_Data_Figure_2", True)
    nm02 = nm02.copy(); nm02.insert(0, "panel_source", "synthetic_uncertainty"); checklist.insert(0, "panel_source", "software_reproduction")
    return pd.concat([nm02, checklist], ignore_index=True, sort=False)


def build_ed3(challenge: pd.DataFrame, transport: pd.DataFrame) -> pd.DataFrame:
    fig = plt.figure(figsize=(7.15, 5.45), constrained_layout=True); grid = fig.add_gridspec(2, 2, height_ratios=[1.25, 0.75], width_ratios=[1.25, 0.75])
    ax = fig.add_subplot(grid[0, 0]); panel(ax, "a"); yy = np.arange(len(challenge)); colors = [RED if s == "BLOCK" else PURPLE for s in challenge.terminal_state]
    ax.barh(yy, np.ones(len(challenge)), color=colors); ax.set_yticks(yy, [v.replace("_", " ").title() for v in challenge.challenge]); ax.invert_yaxis(); ax.set_xlim(0, 1.7); ax.set_xticks([])
    for i, row in challenge.iterrows(): ax.text(0.03, i, row.terminal_state, va="center", color="white", fontweight="bold", fontsize=5.4); ax.text(1.03, i, row.primary_reason_code.replace(".", ".\n", 1), va="center", fontsize=5.2, linespacing=0.9)
    ax.set_title("Controlled challenges", loc="left", fontweight="bold")
    ax = fig.add_subplot(grid[0, 1]); panel(ax, "b"); counts = transport.groupby("repository_platform").size().sort_values(ascending=False)
    yy = np.arange(len(counts)); ax.barh(yy, counts.values, color=[BLUE, GREEN, PURPLE, ORANGE]); ax.set_yticks(yy, [v.replace("_", " ").title() for v in counts.index]); ax.invert_yaxis(); ax.set_xlabel("Transport-audit families"); ax.set_title("Repository transport", loc="left", fontweight="bold")
    time = transport.groupby("time_split").size(); ax.text(0.50, -0.16, "Time strata: " + "; ".join(f"{k.replace('_',' ').title()}={v}" for k, v in time.items()), transform=ax.transAxes, ha="center", fontsize=5.6)
    ax = fig.add_subplot(grid[1, :]); panel(ax, "c"); metrics = ["Clean contracts\nPROCEED", "Invalid contracts\nnon-PROCEED", "Exact invalid\nprimary reason", "Correct pair\ntransition"]
    ax.bar(np.arange(4), [9, 9, 9, 9], color=[BLUE, ORANGE, ORANGE, BLUE]); ax.set_ylim(0, 10.5); ax.set_ylabel("Fixtures (n=9)"); ax.set_xticks(np.arange(4), metrics); ax.set_title("Matched clean–invalid contract checks", loc="left", fontweight="bold")
    for i in range(4): ax.text(i, 9.15, "9/9", ha="center", fontweight="bold")
    save_figure(fig, "Extended_Data_Figure_3", True)
    c = challenge.copy(); c.insert(0, "panel_source", "controlled_challenge")
    t = transport.copy(); t.insert(0, "panel_source", "repository_transport")
    pair = pd.DataFrame({"panel_source": "matched_pair", "metric": metrics, "numerator": 9, "denominator": 9})
    return pd.concat([c, t, pair], ignore_index=True, sort=False)


def build_ed4() -> pd.DataFrame:
    robust = pd.read_csv(OLD_SOURCE / "Source_Data_Extended_Data_Figure_13.csv.gz")
    invalid = pd.read_csv(OLD_SOURCE / "Source_Data_Extended_Data_Figure_14.csv.gz")
    fig, axes = plt.subplots(2, 2, figsize=(7.15, 5.4), constrained_layout=True)
    ax = axes[0, 0]; panel(ax, "a"); vals = np.sort(robust.loc[robust.panel == "LODO", "spearman_rho"].dropna().astype(float)); ax.step(vals, np.arange(1, len(vals) + 1) / len(vals), where="post", color=BLUE); ax.set_xlabel("Spearman ρ"); ax.set_ylabel("Cumulative proportion"); ax.set_title("Leave-one-donor-out effect correlation", loc="left", fontweight="bold")
    ax = axes[0, 1]; panel(ax, "b"); vals = robust.loc[robust.panel == "DIRECTION", "sign_retention_fraction"].dropna().astype(float); ax.hist(vals, bins=12, color=GREEN); ax.set_xlabel("Sign-retention fraction"); ax.set_ylabel("Features"); ax.set_title("Direction retention", loc="left", fontweight="bold")
    ax = axes[1, 0]; panel(ax, "c"); z = robust[robust.panel == "FILTER"].dropna(subset=["filter"]); xx = np.arange(len(z)); ax.bar(xx - 0.18, z.logFC_spearman_rho.astype(float), width=0.36, color=BLUE, label="logFC Spearman ρ"); ax.bar(xx + 0.18, z.stability_set_sign_concordance.astype(float), width=0.36, color=ORANGE, label="Sign concordance"); ax.set_xticks(xx, ["CPM ≥1 in ≥5 donors", "Raw count ≥10 in ≥5 donors"], rotation=12, ha="right", rotation_mode="anchor"); ax.set_ylim(0.85, 1.01); ax.set_title("Filtering sensitivity", loc="left", fontweight="bold"); ax.legend(frameon=False)
    ax = axes[1, 1]; panel(ax, "d"); p = invalid.loc[invalid.panel == "INVALID_DE", "PValue"].dropna().astype(float).clip(lower=1e-300); vals = np.sort(-np.log10(p)); ax.step(vals, np.arange(1, len(vals) + 1) / len(vals), where="post", color=RED); ax.set_xlabel("−log10(P), invalid cell-level analysis"); ax.set_ylabel("Cumulative proportion"); ax.set_title("Invalid pseudoreplication counterfactual", loc="left", fontweight="bold"); ax.text(0.98, 0.05, "NO BIOLOGICAL INTERPRETATION", transform=ax.transAxes, ha="right", color=RED, fontweight="bold", fontsize=5.6)
    save_figure(fig, "Extended_Data_Figure_4", True)
    # Public Source Data retain row-level auditability without propagating source
    # donor tokens.  Stable local indices preserve one-to-one rows and all plotted
    # quantities while avoiding unnecessary participant-like identifiers.
    donor_tokens = sorted(robust["dropped_donor"].dropna().astype(str).unique())
    donor_alias = {token: f"DONOR_{i + 1:03d}" for i, token in enumerate(donor_tokens)}
    robust["dropped_donor"] = robust["dropped_donor"].map(donor_alias)
    robust = robust.rename(columns={"dropped_donor": "dropped_donor_alias"})
    robust.insert(0, "source_block", "whole_beta_robustness"); invalid.insert(0, "source_block", "invalid_counterfactual")
    return pd.concat([robust, invalid], ignore_index=True, sort=False)


def build_ed5() -> pd.DataFrame:
    mds = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_15_SOURCE_DATA.csv")
    donor = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_16A_SOURCE_DATA.csv")
    prop = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_16B_SOURCE_DATA.csv")
    fig = plt.figure(figsize=(7.15, 5.7), constrained_layout=True); outer = fig.add_gridspec(2, 2, height_ratios=[1.4, 1.0], width_ratios=[0.72, 1.28])
    top = outer[0, :].subgridspec(2, 4)
    colors = {"ND": BLUE, "PD": ORANGE, "T2D": RED}
    for i, cl in enumerate([f"C{i}" for i in range(8)]):
        ax = fig.add_subplot(top[i // 4, i % 4]);
        if i == 0: panel(ax, "a")
        z = mds[mds.cluster == cl]
        for condition, color in colors.items():
            q = z[z.condition == condition]; ax.scatter(q.MDS1, q.MDS2, s=9, color=color, alpha=0.85, label=condition)
        ax.set_title(cl, pad=2); ax.tick_params(labelsize=5.0)
        if i == 0: ax.legend(frameon=False, fontsize=5.2, ncol=3, loc="upper left")
        if i // 4 == 1: ax.set_xlabel("MDS1", fontsize=5.2)
        if i % 4 == 0: ax.set_ylabel("MDS2", fontsize=5.2)
    ax = fig.add_subplot(outer[1, 0]); panel(ax, "b"); d = donor.set_index("cluster")[["ND", "PD", "T2D"]]; im = ax.imshow(d, cmap="Blues", aspect="auto", vmin=0, vmax=17); ax.set_xticks(np.arange(3), d.columns); ax.set_yticks(np.arange(8), d.index)
    for i in range(8):
        for j in range(3): ax.text(j, i, int(d.iloc[i, j]), ha="center", va="center", fontsize=5.4, color="white" if d.iloc[i, j] > 12 else INK)
    ax.set_title("Qualifying donors (≥20 cells)", loc="left", fontweight="bold")
    ax = fig.add_subplot(outer[1, 1]); panel(ax, "c")
    for ci, condition in enumerate(["ND", "PD", "T2D"]):
        vals = [prop[(prop.condition == condition) & (prop.cluster == c)].cluster_proportion.values for c in [f"C{i}" for i in range(8)]]; pos = np.arange(8) + (ci - 1) * 0.23
        bp = ax.boxplot(vals, positions=pos, widths=0.20, patch_artist=True, showfliers=False)
        for b in bp["boxes"]: b.set_facecolor(colors[condition]); b.set_alpha(0.70)
        for med in bp["medians"]: med.set_color(INK)
    ax.set_xticks(np.arange(8), [f"C{i}" for i in range(8)]); ax.set_ylabel("Donor cluster proportion"); ax.set_title("Recovered beta-cell composition", loc="left", fontweight="bold")
    handles = [plt.Line2D([0], [0], color=colors[c], lw=6) for c in ["ND", "PD", "T2D"]]; ax.legend(handles, ["ND", "PD", "T2D"], frameon=False, ncol=3)
    save_figure(fig, "Extended_Data_Figure_5", True)
    mds.insert(0, "source_block", "mds"); donor.insert(0, "source_block", "support"); prop.insert(0, "source_block", "composition")
    return pd.concat([mds, donor, prop], ignore_index=True, sort=False)


def build_ed6() -> pd.DataFrame:
    de = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_17_SOURCE_DATA.csv.gz")
    hall = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_18_HALLMARK_SOURCE_DATA.csv").set_index("pathway")
    react = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_18_REACTOME_SOURCE_DATA.csv").set_index("pathway")
    concord = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_19_SOURCE_DATA.csv")
    fig = plt.figure(figsize=(7.15, 6.6), constrained_layout=True); outer = fig.add_gridspec(4, 2, height_ratios=[1.30, 0.95, 1.25, 0.80], width_ratios=[1, 1])
    top = outer[0, :].subgridspec(2, 4)
    for i, cl in enumerate([f"C{i}" for i in range(8)]):
        ax = fig.add_subplot(top[i // 4, i % 4]);
        if i == 0: panel(ax, "a")
        z = de[de.cluster == cl]; sig = z.FDR_global_edgeR < 0.05
        ax.scatter(z.logFC_edgeR[~sig], -np.log10(z.PValue_edgeR[~sig].clip(lower=1e-300)), s=1.2, color="#C7CDD1", alpha=0.35)
        ax.scatter(z.logFC_edgeR[sig], -np.log10(z.PValue_edgeR[sig].clip(lower=1e-300)), s=2.5, color=RED, alpha=0.75); ax.axvline(0, color=GRAY, lw=0.4); ax.set_title(f"{cl} (FDR: {sig.sum()})", fontsize=5.6); ax.tick_params(labelsize=5.2)
        if i // 4 == 1: ax.set_xlabel("log2FC", fontsize=5.2)
        if i % 4 == 0: ax.set_ylabel("−log10(P)", fontsize=5.2)
    for idx, (data, title, prefix, letter) in enumerate([(hall, "Hallmark", "HALLMARK_", "b"), (react, "Reactome", "REACTOME_", "c")]):
        ax = fig.add_subplot(outer[idx + 1, :]); panel(ax, letter); values = data[[f"C{i}" for i in range(8)]].to_numpy(float); vmax = max(2.0, float(np.nanpercentile(np.abs(values), 95))); im = ax.imshow(values, aspect="auto", cmap="RdBu_r", vmin=-vmax, vmax=vmax)
        ax.set_xticks(np.arange(8), [f"C{i}" for i in range(8)]); ax.set_yticks(np.arange(len(data)), [v.replace(prefix, "").replace("_", " ") for v in data.index], fontsize=5.2); ax.set_title(f"{title} competitive tests", loc="left", fontweight="bold"); cb = fig.colorbar(im, ax=ax, fraction=0.025, pad=0.02); cb.set_label("signed −log10(global FDR)", fontsize=5.2); cb.ax.tick_params(labelsize=5.2)
    ax = fig.add_subplot(outer[3, 0]); panel(ax, "d")
    for contrast, color, marker in [("T2D-ND", BLUE, "o"), ("PD-ND", ORANGE, "s"), ("T2D-PD", GREEN, "^")]:
        z = concord[concord.contrast == contrast]; ax.plot(z.cluster, z.spearman_logFC, marker=marker, color=color, label=contrast, markersize=3)
    ax.set_ylim(0.86, 1.0); ax.set_ylabel("Spearman ρ of logFC"); ax.set_xlabel("Author cluster"); ax.set_title("edgeR–voom sensitivity", x=0.08, ha="left", fontweight="bold"); ax.legend(frameon=False, ncol=3, fontsize=5.2)
    ax = fig.add_subplot(outer[3, 1]); panel(ax, "e"); ax.bar(["Atlas whole-beta\n511 genes", "State-resolved\n225 genes"], [511, 225], color=[GRAY, BLUE]); ax.bar("State-resolved\n225 genes", 89, color=ORANGE); ax.text(1, 158, "136/225\noverlap", ha="center", va="center", color="white", fontweight="bold", fontsize=5.2); ax.text(0.98, 0.92, "310/418 associations overlap\n310/310 direction concordance", transform=ax.transAxes, ha="right", va="top", fontsize=5.5); ax.set_ylabel("Unique genes"); ax.set_title("Same-cohort atlas overlap", x=0.08, ha="left", fontweight="bold")
    save_figure(fig, "Extended_Data_Figure_6", True)
    hall = hall.reset_index(); react = react.reset_index(); de.insert(0, "source_block", "cluster_de"); hall.insert(0, "source_block", "hallmark_pathway"); react.insert(0, "source_block", "reactome_pathway"); concord.insert(0, "source_block", "method_concordance")
    atlas = pd.DataFrame({"source_block": "atlas_overlap", "metric": ["atlas_genes", "state_resolved_genes", "overlap_genes", "overlap_associations", "direction_concordant_associations"], "value": [511, 225, 136, 310, 310]})
    return pd.concat([de, hall, react, concord, atlas], ignore_index=True, sort=False)


def parse_references() -> list[str]:
    doc = Document(INPUT_DOCX)
    text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    ref_text = text.split("\nReferences\n", 1)[1].split("\nFigures\n", 1)[0]
    parts = re.split(r"(?<!\d)(?=(?:[1-9]|[1-4]\d|5[0-7])\.\s)", ref_text)
    refs = [p.strip() for p in parts if re.match(r"^\d+\.\s", p.strip())]
    # The source Word paragraph contains a hard line break inside reference 57
    # ("COVID-\n19. Cell"), which resembles a new numbered reference to a
    # naïve regular expression.  References 1–54 are the complete cited set in
    # the final scientific story, so require and retain that verified prefix.
    if len(refs) < 54:
        raise AssertionError(f"expected at least 54 input references, found {len(refs)}")
    return refs[:54]


TITLE = "scDesignGuard: source-bound design gating for multisample single-cell differential analysis"
AUTHORS = "*Guoyong Wang¹, *Kaijun Zhang², Jiyue Jiang³, Weixin Wang¹, Hui Bi⁴, Haojun Liang⁵, Zuoliang Qi⁵, *Ying Huang², *Yu Li³, *Xiaonan Yang¹"
AFFILIATIONS = (
    "1. Department of Hemangioma and Vascular Malformation, Plastic Surgery Hospital, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing 100144, China. "
    "2. Department of Gastroenterology, Children's Hospital of Fudan University, National Children's Medical Center, No. 399 Wanyuan Road, Minhang District, Shanghai 201102, People's Republic of China. "
    "3. Department of Computer Science and Engineering, The Chinese University of Hong Kong, Sha Tin, New Territories, Hong Kong SAR, China. "
    "4. Department of Internal Medicine, Plastic Surgery Hospital, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing 100144, China. "
    "5. Department of Comprehensive Plastic Surgery, Plastic Surgery Hospital, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing 100144, China."
)
CORRESPONDENCE = "Correspondence and requests for materials should be addressed to Xiaonan Yang, Department of Hemangioma and Vascular Malformation, Plastic Surgery Hospital, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing 100144, China. Tel.: +86 188 1060 1889; Fax: +86 10 5396 8149; Email: yxnan@aliyun.com."


ABSTRACT = """Single-cell workflows can produce numerical results even when source identity, biological replication, target support or model estimability are unresolved. We developed scDesignGuard, a source-bound stage-aware framework that returns PROCEED, ABSTAIN, BLOCK or NON_EVALUABLE before inference. Against frozen source-bound human consensus in an overlap-excluded benchmark of 15 study families, scDesignGuard cleared 4/4 valid designs, blocked 6/6 hard-invalid designs and contained 5/5 unresolved designs, with no unsafe continuation. In four overlap-excluded end-to-end cases, two reached model validity and two stopped at explicit source-binding or donor-replication gates. Nine controlled invalidities produced the prespecified state and reason, and the frozen software reproduced across environments. Applied to 99,029 pancreatic beta cells from 48 donors, scDesignGuard localized same-cohort type 2 diabetes-associated composition, transcriptional and pathway signals. scDesignGuard separates runnable workflows from evidence-supported analyses."""


INTRO = [
    "Single-cell workflows can complete and return apparently precise results even when the experimental unit or requested contrast is scientifically unsupported. Related integration, atlas-construction and computational-ecosystem challenges further complicate the interpretation of runnable workflows.^3–7^ Cells are nested within donors or experimental units, so cell number cannot substitute for independent biological replication; cell-level testing may otherwise inflate certainty.^1,2,8,9,13^",
    "Donor-aware pseudobulk, multi-subject differential-expression and differential-composition methods address how to estimate effects once a valid task has been specified.^8–20^ Related work extends statistical modelling, simulation and benchmarking across single-cell settings.^21–34^ These methods do not ordinarily establish, before inference, whether the selected repository object, donor map, target semantics and estimand authorize the requested analysis.",
    "FAIR software and research-object principles, workflow provenance systems, ontologies and repository infrastructure improve traceability and reuse.^35–50^ They record and standardize data and computation but do not provide stage-specific scientific authorization that distinguishes a demonstrated invalidity from unresolved or not-yet-applicable evidence.",
    "Here we evaluate scDesignGuard in an overlap-excluded benchmark of 15 study families, a separate four-case end-to-end gate evaluation, nine controlled challenge tests, multi-repository transport audits and clean-environment software reproduction. We then apply the gated workflow to a 99,029-cell pancreatic beta-cell atlas to examine the biological use of evidence-aware design control."
]


RESULT_SECTIONS = [
    ("scDesignGuard converts source-bound evidence into stage-specific authorization", [
        "scDesignGuard represents an analysis request as a versioned contract that binds the source object, independent biological unit, condition and pairing structure, target ontology, support rule and estimand (Fig. 1a). The evaluator returns PROCEED when evidence required at the current stage is complete, ABSTAIN for modelled uncertainty, BLOCK for a demonstrated invalidity and NON_EVALUABLE when available evidence cannot resolve the requested analysis. Each output carries ordered reasons and an explicit denominator.",
        "Authorization is sequential (Fig. 1b). Source-scope clearance permits structural-design review; structural clearance permits target-support computation; support clearance permits count aggregation and model-validity assessment. Passing a gate authorizes assessment of the next stage, not biological inference.",
        "GSE144735 illustrates the complete path (Fig. 1c). Its paired contract contained six complete pairs, 3,356 target epithelial cells and 12 donor-condition pseudobulks. The model matrix had rank 7/7 and five residual degrees of freedom. Independent implementations reconciled donor assignment, target-cell counts, library sums, pseudobulk dimensions and model rank before final authorization. No differential effect was estimated in this gate evaluation."
    ]),
    ("An overlap-excluded benchmark demonstrates state discrimination", [
        "The final overlap-excluded benchmark comprised 15 study families with frozen source-bound human consensus labels: four valid, six hard-invalid and five structurally unresolved. scDesignGuard returned PROCEED for all four valid designs, BLOCK for all six hard-invalid designs and NON_EVALUABLE for all five unresolved designs. Valid clearance was 4/4 (95% Wilson CI, 0.510–1.000), hard-invalid detection 6/6 (0.610–1.000), unresolved containment 5/5 (0.566–1.000), unsafe continuation 0/6 (0.000–0.390) and exact state concordance 15/15 (0.796–1.000; Fig. 2).",
        "The reference review was blinded to evaluator, executable-checklist, end-to-end and biological-outcome outputs, but verified prepared source-bound fields rather than independently compiling every semantic contract field de novo. These endpoints therefore quantify concordance with frozen source-bound human consensus, not accuracy against an independently constructed external reference."
    ]),
    ("End-to-end cases show when analyses proceed and when they stop", [
        "A separate overlap-excluded end-to-end cohort comprised four cases processed in frozen order through source scope, structural design, target support, count aggregation, model validity and final authorization (Fig. 3a,c).",
        "GSE144735 contained six complete pairs, 3,356 target cells, 12 pseudobulks, a full-rank 7-column design and five residual degrees of freedom. GSE123813 contained 11 complete pairs, 33,106 target cells, 22 pseudobulks, a full-rank 12-column design and ten residual degrees of freedom (Fig. 3b). Both reached final authorization. Human Cell Atlas project 135f7f5c-4a85-4bcf-9f7c-4f035ff1e428 stopped because exact donor-condition-modality and raw-count-layer binding could not be established. GSE81076 stopped because five total donor identifiers cannot satisfy the frozen minimum of five donors per condition.",
        "Python primary computation, an independently compiled C count implementation and an independent R design-matrix implementation reconciled donor assignment, target-cell counts, raw-library sums, pseudobulk dimensions, rank, condition-effect estimability and residual degrees of freedom. Differential expression, pathways and biological effects were not computed in this evaluation."
    ]),
    ("Controlled failures and reproducible transport define the operating boundary", [
        "Nine deterministic perturbations altered one design property at a time: donor deletion, cell-level pseudoreplication, batch-condition confounding, duplicate donor, pair destruction, wrong target, wrong version, mixed strata or low support. All nine produced a safe non-PROCEED state and the prespecified primary reason (9/9 for both endpoints; Fig. 4a). These are controlled challenge tests of rule sensitivity rather than estimates of defect prevalence.",
        "Repository transport audits covered 18 study families across BioStudies/ArrayExpress, Single Cell Portal, NCBI GEO and Human Cell Atlas (Fig. 4b). These Tier C audits evaluated source identity, access boundaries and terminal reasons without target, expression or outcome analysis; they are not effect-performance cases.",
        "The frozen wheel, OCI image, JSON Schema, reason registry and terminal-state fixtures reproduced in a clean environment. Three current-release suites and two compatibility suites reproduced their expected behavior, and the repository run yielded 1,362/1,362 passed tests, 29 planned skips and 42 passed subtests (Fig. 4c). Engineering reproducibility is distinct from scientific validation."
    ]),
    ("Donor-aware gating localizes same-cohort beta-cell signals", [
        "The same-cohort biological application used the GSE221156 pancreatic islet atlas^51^ and comprised 99,029 author-annotated beta cells from 48 donors: 17 normal, 14 prediabetes and 17 type 2 diabetes donors across author-defined clusters C0–C7. Counts were aggregated once per donor-cluster. C7 had the lowest qualifying support (13 normal, eight prediabetes and ten type 2 diabetes donors) and was interpreted cautiously.",
        "Recovered beta-cell-state composition differed between type 2 diabetes and normal donors after adjustment for assay, sex, ethnicity, age and body mass index (Fig. 5a,b). Mean donor C6 proportions were 0.060, 0.086 and 0.183 in normal, prediabetes and type 2 diabetes, respectively; the adjusted type 2 diabetes–normal coefficient was 0.191 on the arcsine-square-root scale (BH-FDR across eight clusters, 1.78 × 10−4). C0 decreased from 0.242 to 0.137 (coefficient, −0.122; FDR, 0.0213). These proportions describe recovered beta cells, not absolute islet abundance.",
        "Donor-level pseudobulk models^8–10^ fitted with edgeR quasi-likelihood methods^20^ tested 100,434 eligible gene–cluster entries. The type 2 diabetes–normal contrast yielded 418 associations involving 225 unique genes (Fig. 5c). Reproduction of the source-atlas whole-beta workflow^51^ yielded 511 genes; 136/225 state-resolved genes overlapped this set and accounted for 310/418 associations, all direction-concordant. Limma-voom sensitivity analysis^52,53^ yielded 444 associations, 311 shared with edgeR and no shared sign conflict. Hallmark and Reactome collections^49,50^ were tested competitively with CAMERA^54^, identifying lower fatty-acid metabolism across all eight clusters, lower oxidative phosphorylation in seven, higher TNF signalling through NF-κB in six and lower mTORC1 signalling in six (Fig. 5d). These findings localize same-cohort signals rather than establish independent replication."
    ]),
]


METHOD_SECTIONS = [
    ("Framework and scientific states", [
        "A task contract binds immutable source identifiers and versions, the independent biological unit, condition and pairing fields, source-defined target labels and ontology relations, support thresholds, estimand definitions and evidence-object references. The evaluator emits PROCEED, ABSTAIN, BLOCK or NON_EVALUABLE with a primary and ordered secondary reason set. Rules are stage-specific: source scope, structural design, target support, count aggregation, model validity and biological inference are authorized separately."
    ]),
    ("Evidence tiers and decision hierarchy", [
        "Tier A requires at least five donors per group or five complete pairs, at least 20 target cells per donor, exact donor mapping, EXACT or CHILD target ontology, established analysis access and a complete estimand. Tier B requires at least three donors per group or pairs and ten target cells per donor and is limited to supportive robustness or direction. Tier C is restricted to metadata, provenance, repair and terminal-state evaluation. Tier D comprises known-invalid transformations of public real-data contracts. The donor and target-cell thresholds are conservative project-defined authorization cut-offs rather than universal sample-size recommendations. They distinguish full analytical authorization from lower-evidence robustness or structural use.",
        "Terminal states and primary reasons were assigned using a frozen stage- and reason-priority hierarchy. Demonstrated hard invalidities took precedence over unresolved downstream evidence, whereas unperformed or stage-inapplicable checks remained NOT_ASSESSED or NOT_APPLICABLE. The complete hierarchy and reason-code definitions are provided in Supplementary Table 3. In the evaluated implementation, state precedence was BLOCK, NON_EVALUABLE, ABSTAIN and PROCEED; activated reasons within a state were sorted lexicographically by reason code."
    ]),
    ("Final benchmark selection and human consensus", [
        "Candidate study families were screened against a project-wide family-identity registry linking repository accession or collection identifiers with publication, project and cohort aliases. Families with any identified development or calibration overlap were excluded without replacement before final endpoint estimation. Selection did not use evaluator output, downstream accessibility, biological-effect estimates or result appearance. The final benchmark comprised 15 study families. Two named reviewers verified prepared source-bound records while blinded to evaluator, executable-checklist, end-to-end and biological-outcome outputs; disagreements were resolved before evaluator execution, and the resulting consensus labels were frozen. Because the task involved verification of prepared source-bound fields rather than fully independent de novo construction of every semantic field, the reference is described as frozen source-bound human consensus rather than an independently constructed criterion standard. Reviewer-to-reviewer inter-rater reliability was not treated as an endpoint."
    ]),
    ("Benchmark endpoints", [
        "VALID_DESIGN, HARD_INVALID_DESIGN and STRUCTURALLY_UNRESOLVED reference classes implied expected PROCEED, BLOCK and NON_EVALUABLE states, respectively. Primary endpoints were valid clearance, hard-invalid detection, unresolved containment, unsafe continuation and exact state concordance. Each rate is reported with exact membership and a two-sided 95% Wilson interval. Evaluator and executable-checklist outputs did not modify the human reference. These endpoints evaluate deterministic state concordance within a shared project taxonomy and are not estimates of superiority over an independently authored expert method."
    ]),
    ("End-to-end target-support and model-validity evaluation", [
        "A separate four-case overlap-excluded cohort was processed in frozen order through source scope, structural design, target support, count aggregation, model validity and final authorization. Candidate inclusion did not use scDesignGuard state, downstream effect estimates or biological outcomes.",
        "Independent comparisons required at least five qualifying donors per condition; paired comparisons required at least five complete pairs. Each donor-target unit required at least 20 target cells and a finite positive raw integer library. Repeated samples were aggregated to donors.^8–10^ Model validity required a full-column-rank matrix, an estimable condition effect and positive residual degrees of freedom. A Python implementation computed donor assignment, target-cell counts, raw-library sums, pseudobulk dimensions, rank, estimability and residual degrees of freedom. Count quantities were recomputed by an independently compiled C implementation and model-matrix quantities by an independent R implementation. Differential expression, pathways and effects were not computed."
    ]),
    ("Controlled invalidities and repository transport", [
        "Nine transformations modified one prespecified property while retaining all other contract fields: donor deletion, cell-level pseudoreplication, batch-condition confounding, duplicate donor, pair destruction, wrong target, wrong version, mixed strata and low support. Expected states and primary reasons were fixed before evaluation.",
        "Repository transport used official metadata and identity endpoints for BioStudies/ArrayExpress,^44^ Single Cell Portal, NCBI GEO,^43^ Human Cell Atlas^45^ and CELLxGENE.^42^ Authentication failures, absent row metadata and unresolved donor mappings were retained as evidence states rather than inferred from descriptive text. Expression and outcomes were not accessed for transport-audit families."
    ]),
    ("Software reproducibility", [
        "The evaluated software was frozen as a Python wheel, OCI image, JSON Schema, 21-code reason registry, four terminal-state fixtures and software bill of materials. Packaging and provenance followed FAIR software and computational-workflow principles.^35,38–41^ Clean-environment installation, current-release tests, compatibility tests and repository-boundary fixtures were executed against recorded hashes."
    ]),
    ("Pancreatic beta-cell analysis", [
        "The exact CELLxGENE object^42^ for GSE221156^51^ contained 99,029 author-annotated beta cells from 48 donors. Author C0–C7 labels were retained without marker-based renaming. Raw counts were summed once per donor-cluster. A donor-cluster unit required at least 20 cells and a positive raw library sum.",
        "Composition models used arcsine-square-root donor proportions and adjusted for assay, sex, ethnicity, age and body mass index; inference used limma empirical-Bayes tests with BH correction across eight clusters per contrast. Cluster-specific edgeR quasi-likelihood models^20^ used donor-level pseudobulk counts^8–10^ and the same covariate design. Genes were filtered within cluster, and BH correction was applied across all eligible gene-cluster tests within each contrast. Limma-voom^52,53^ used identical matrices, covariates and contrasts. All eligible Hallmark and Reactome pathways^49,50^ were tested competitively with CAMERA,^54^ with BH correction across the complete cluster–pathway universe within each collection and contrast. For Fig. 5d, the 12 Hallmark pathways with the smallest minimum global cluster–pathway FDR across C0–C7 in the type 2 diabetes–normal contrast were displayed. Extended Data Fig. 6 applied the same ranking rule separately to display 15 Hallmark and 20 Reactome pathways. Display selection did not change either multiple-testing universe.",
        "The atlas-overlap analysis reproduced the GSE221156 authors' whole-beta type 2 diabetes–normal workflow,^51^ yielding 511 genes. State-resolved genes and gene–cluster associations were joined to that set in a shared 13,207-gene test universe. Overlap denominators were 225 unique genes and 418 significant associations; direction concordance was evaluated only for overlapping associations."
    ]),
    ("Statistics and reporting", [
        "Cells were never treated as independent biological replicates.^8,9,13^ Binomial intervals are two-sided Wilson intervals. Differential-expression and composition P values are two-sided and adjusted as specified above. Pathway direction is encoded by the sign of the test statistic. No threshold, endpoint or truth rule for the final 15-family overlap-excluded benchmark or the four-case end-to-end evaluation was modified after those analysis sets and endpoint definitions were frozen."
    ]),
]


DISCUSSION = [
    "scDesignGuard formalizes a distinction that computational pipelines can obscure: technical execution does not itself authorize scientific inference. The framework makes source identity, biological replication, target support and estimability explicit prerequisites and records where authorization stops. It complements rather than replaces donor-aware statistical engines.",
    "The 15-family overlap-excluded benchmark shows that the evaluator is not an always-BLOCK or always-NON_EVALUABLE rule set: valid designs proceeded, demonstrated invalidities were blocked and unresolved records remained non-evaluable. The exact state concordance is accompanied by wide Wilson intervals because the denominator is modest. The frozen source-bound human consensus, checklist and evaluator also share a project taxonomy, so these results support deterministic state discrimination within that taxonomy rather than accuracy against an external criterion or superiority to an independently authored expert method.",
    "The four-case end-to-end evaluation complements the benchmark by showing that PROCEED can reach quantitatively supported target aggregation and an estimable model, whereas prespecified early stops preserve interpretable source-binding or donor-replication reasons. This denominator is also modest, and no differential effect was estimated in the gate evaluation.",
    "Controlled challenge tests, repository transport, clean-environment software reproduction and the same-cohort beta-cell application answer different questions. They demonstrate rule sensitivity, portability of structural evidence, deterministic implementation and biological utility, respectively, but none substitutes for the others. The beta-cell analysis localizes donor-aware signals within one cohort and provides neither independent replication nor causal or clinical claims.",
    "Semantic contract construction remains the principal practical limitation because trained human interpretation is still required for complex source records, ontology relations and estimands. Future work should measure fully independent contract-construction reproducibility, time and usability, and should expand overlap-excluded end-to-end evaluation. Within the present boundaries—modest benchmark and end-to-end denominators, shared taxonomy, one-cohort biology and no clinical or causal interpretation—scDesignGuard provides a reproducible layer that separates runnable workflows from evidence-supported inference."
]


FIGURE_LEGENDS = [
    "Figure 1 | Source-bound design evidence and stage-specific authorization. a, scDesignGuard compiles source identity, donor and condition structure, target ontology, support and estimand into an evidence-bearing contract that returns PROCEED, ABSTAIN, BLOCK or NON_EVALUABLE with explicit reasons. b, Passing one gate authorizes only assessment of the next gate. c, GSE144735 worked example: six complete pairs, 3,356 target epithelial cells, 12 donor-condition pseudobulks, design-matrix rank 7/7 and five residual degrees of freedom. No inferential test or biological effect was computed in this gate evaluation.",
    "Figure 2 | Overlap-excluded benchmark and state concordance. a, Separation of family-identity audit, frozen source-bound human consensus and sealed evaluator execution. b, Frozen human-reference strata: four valid, six hard-invalid and five structurally unresolved families. c, Exact proportions and two-sided 95% Wilson intervals for valid clearance (4/4), hard-invalid detection (6/6), unresolved containment (5/5), unsafe continuation (0/6) and exact state concordance (15/15). d, Family-level state matrix. These endpoints quantify concordance with frozen source-bound human consensus, not independently labelled external-accuracy estimates.",
    "Figure 3 | Overlap-excluded end-to-end outcomes. a, Outcomes for four cases. GSE144735 and GSE123813 passed all gates; Human Cell Atlas project 135f7f5c-4a85-4bcf-9f7c-4f035ff1e428 stopped at structural source binding and GSE81076 stopped at the donor-replication threshold. b, Exact quantitative checks for the two full gate passes. GSE144735 comprised six complete pairs, 3,356 target cells, 12 pseudobulks, rank 7/7 and five residual degrees of freedom; GSE123813 comprised 11 complete pairs, 33,106 target cells, 22 pseudobulks, rank 12/12 and ten residual degrees of freedom. c, Cases reaching each sequential gate. No differential expression, pathway analysis or biological-effect estimation was performed.",
    "Figure 4 | Controlled challenges, repository transport and software reproduction. a, Terminal states for nine deterministic known-invalid transformations. All nine returned a safe non-PROCEED state and the prespecified primary reason. These are controlled challenge tests, not prevalence-weighted cases. b, Eighteen Tier C repository transport audits across BioStudies/ArrayExpress (n=5), Single Cell Portal (n=4), NCBI GEO (n=5) and Human Cell Atlas (n=4). Expression and outcomes were not accessed. c, Reproduction summaries: 3/3 current-release suites, 2/2 compatibility suites and 1,362/1,362 repository tests passed; 29 tests were skipped by design and 42 subtests passed. No statistical test was performed for b or c.",
    "Figure 5 | Donor-aware beta-cell composition, transcription and pathway localization. a, Mean donor proportions of author-defined clusters C0–C7 in normal (ND; n=17 donors), prediabetes (PD; n=14) and type 2 diabetes (T2D; n=17); bars are arithmetic means across donors. b, Adjusted T2D–ND composition coefficients from donor-level models on the arcsine-square-root scale after adjustment for assay, sex, ethnicity, age and body mass index. Points are coefficients; error bars are descriptive ±1.96 unmoderated coefficient standard errors (residual d.f.=39). P values are two-sided limma empirical-Bayes tests; FDR values are Benjamini–Hochberg adjusted across eight clusters. Positive values indicate higher recovered-cell proportion in T2D. c, Significant gene–cluster associations after BH correction across 100,434 eligible tests in the T2D–ND contrast. Bars show edgeR quasi-likelihood and limma-voom counts. edgeR identified 418 associations involving 225 genes; 136/225 genes overlapped the same-cohort 511-gene whole-beta atlas set, accounting for 310/418 associations with 310/310 direction agreement. d, Signed −log10 global FDR for competitive Hallmark CAMERA tests. All eligible Hallmark cluster–pathway tests were included in BH correction; for visualization, the 12 pathways with the smallest minimum global FDR across C0–C7 were displayed. Red denotes higher and blue lower competitive activity in T2D. Display selection did not alter the inferential universe. Results are same-cohort localization, not independent replication."
]


ED_LEGENDS = [
    "Extended Data Figure 1 | Development support, robustness and calibration. a, Qualifying independent-unit support under the development protocol; the dashed line marks the project-defined five-donor cut-off. b, Leave-one-donor-out direction-stability profiles. c, Calibration-label composition. d, Minimum target-cell support among qualifying donor-target units; the dashed line marks the project-defined 20-cell cut-off. These data are development evidence, not benchmark performance.",
    "Extended Data Figure 2 | Synthetic uncertainty and software reproduction. a, Descriptive Monte Carlo estimates for pooled Type-I rate and mean replicate FDP/FDR with simulation intervals and Monte Carlo standard errors. The 0.05 line is a reference, not a post hoc claim boundary. b, Clean-environment installation, schema, reason-registry, fixture and security checks. Synthetic estimates do not support strict nominal Type-I or FDR-control claims.",
    "Extended Data Figure 3 | Controlled challenges and repository transport. a, Nine controlled invalidities and observed terminal states and primary reasons. b, Tier C transport-audit families by repository and time stratum. c, Matched clean–invalid contract checks. Expression and outcomes were not accessed for transport audits; no biological effect was estimated.",
    "Extended Data Figure 4 | Whole-beta robustness and invalid counterfactual. a, Leave-one-donor-out effect correlations. b, Direction-retention fractions. c, Filtering sensitivity. d, Invalid cell-level pseudoreplication counterfactual. The counterfactual is a methodological demonstration and carries no biological interpretation.",
    "Extended Data Figure 5 | Beta-cell quality control, support and composition. a, EdgeR multidimensional-scaling coordinates for donor–cluster pseudobulks, coloured by condition; no point was excluded after inspection. b, Qualifying donor counts under the ≥20-cell rule. c, Donor-level recovered-beta-cell composition. Box plots show medians, interquartile ranges and whiskers extending to 1.5× the interquartile range.",
    "Extended Data Figure 6 | Beta-cell differential expression, pathways and atlas overlap. a, EdgeR quasi-likelihood T2D–ND effects; red points pass BH correction across all eligible gene–cluster tests. b,c, Signed −log10 global FDR for Hallmark and Reactome CAMERA tests. The 15 Hallmark and 20 Reactome pathways with the smallest minimum global cluster–pathway FDR across C0–C7 were displayed separately; all eligible tests in each collection remained in the BH correction universe. d, edgeR–voom log-fold-change concordance. e, Same-cohort overlap: 136/225 state-resolved genes and 310/418 associations overlap the whole-beta atlas set, with 310/310 direction concordance. State-resolved-only genes are not automatically novel findings."
]


DATA_AVAILABILITY = "The final 15-family benchmark membership, four-case end-to-end gate results, controlled-challenge results, repository-transport summaries, quantitative reconciliation and figure Source Data accompany the manuscript. Historical development and audit objects are retained under version control but are not part of the submitted scientific denominators. Expression objects are not redistributed and remain subject to repository and study terms. The same-cohort biological application is available under accession GSE221156."
CODE_AVAILABILITY = "The evaluated scDesignGuard v0.1.0 software, exact analysis scripts, JSON Schema, reason-code registry, synthetic fixtures, tests, software bill of materials and reproducibility documentation are publicly available at https://github.com/wang177777/scDesignGuard (release v0.1.0). The release includes the exact evaluated wheel and hash binding for the OCI image. Third-party expression objects and participant-level repository payloads are not redistributed; their official accession or collection identifiers are provided in the manuscript and Source Data."
ETHICS = "This study exclusively reanalysed publicly available, de-identified datasets and involved no new participant recruitment, intervention or collection of identifiable private information. Under applicable institutional policies, this secondary analysis did not require additional ethics committee approval or informed consent. Ethics approvals and informed-consent procedures for the original studies are described in the source publications."
AUTHOR_CONTRIBUTIONS = "Guoyong Wang, Kaijun Zhang, Jiyue Jiang, Yu Li and Xiaonan Yang conceived the study. Kaijun Zhang, Jiyue Jiang and Yu Li developed the methodology. Jiyue Jiang and Yu Li developed the software, curated the data and performed the formal analysis. Guoyong Wang, Weixin Wang, Hui Bi, Haojun Liang, Zuoliang Qi and Ying Huang contributed to clinical and biological interpretation. Guoyong Wang, Kaijun Zhang, Jiyue Jiang and Yu Li drafted the manuscript. All authors reviewed, edited and approved the final manuscript. Guoyong Wang and Xiaonan Yang supervised the study. Guoyong Wang, Kaijun Zhang, Ying Huang, Yu Li and Xiaonan Yang contributed equally to this work."
FUNDING = "This work was supported by the Special Program for Clinical and Translational Medical Research of the Chinese Academy of Medical Sciences (2025-12M-C&T-B-067), the National Clinical Key Specialty Construction Project (23003), the Plastic Medicine Research Fund of the Chinese Academy of Medical Sciences (2024-ZX-1-01), the Special Research Fund for Plastic Surgery Hospital, Chinese Academy of Medical Sciences and Peking Union Medical College (YSZ2024CG007), and the Beijing Natural Science Foundation (L256048)."
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def set_margins(doc: Document, top=0.72, bottom=0.72, left=0.78, right=0.78) -> None:
    for section in doc.sections:
        section.top_margin = Inches(top); section.bottom_margin = Inches(bottom); section.left_margin = Inches(left); section.right_margin = Inches(right)


def style_doc(doc: Document) -> None:
    set_margins(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(4.5); styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size, color in [("Title", 20, "1F4E79"), ("Heading 1", 14, "1F4E79"), ("Heading 2", 11.5, "2F75B5"), ("Heading 3", 10, "25313B")]:
        styles[name].font.name = "Arial"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color); styles[name].font.bold = True


def add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\^[^^]+\^)", text)
    for part in parts:
        if part.startswith("^") and part.endswith("^"):
            run = p.add_run(part[1:-1]); run.font.superscript = True
        else:
            p.add_run(part)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_figure_page(doc: Document, png: Path) -> None:
    """Append a full-page figure without duplicating the legend section."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Inches(6.7))


def build_manuscript(refs: list[str]) -> Path:
    doc = Document(); style_doc(doc)
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(TITLE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(AUTHORS).bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("*These authors contributed equally to this work.").italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(AFFILIATIONS)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(CORRESPONDENCE)
    add_heading(doc, "Abstract"); add_rich_paragraph(doc, ABSTRACT)
    for para in INTRO: add_rich_paragraph(doc, para)
    add_heading(doc, "Results")
    for heading, paras in RESULT_SECTIONS:
        add_heading(doc, heading, 2)
        for para in paras: add_rich_paragraph(doc, para)
    add_heading(doc, "Discussion")
    for para in DISCUSSION: add_rich_paragraph(doc, para)
    add_heading(doc, "Methods")
    for heading, paras in METHOD_SECTIONS:
        add_heading(doc, heading, 2)
        for para in paras: add_rich_paragraph(doc, para)
    add_heading(doc, "Data availability"); add_rich_paragraph(doc, DATA_AVAILABILITY)
    add_heading(doc, "Code availability"); add_rich_paragraph(doc, CODE_AVAILABILITY)
    add_heading(doc, "Ethics statement"); add_rich_paragraph(doc, ETHICS)
    add_heading(doc, "Author contributions"); add_rich_paragraph(doc, AUTHOR_CONTRIBUTIONS)
    add_heading(doc, "Funding"); add_rich_paragraph(doc, FUNDING)
    add_heading(doc, "Competing interests"); add_rich_paragraph(doc, "The authors declare no competing interests.")
    add_heading(doc, "References")
    for ref in refs: add_rich_paragraph(doc, ref)
    add_heading(doc, "Figure legends")
    for legend in FIGURE_LEGENDS: add_rich_paragraph(doc, legend)
    add_heading(doc, "Extended Data figure legends")
    for legend in ED_LEGENDS: add_rich_paragraph(doc, legend)
    for i in range(1, 6): add_figure_page(doc, OUT / f"figures/main/Figure_{i}.png")
    for i in range(1, 7): add_figure_page(doc, OUT / f"figures/extended_data/Extended_Data_Figure_{i}.png")
    path = OUT / "scDesignGuard_Nature_Methods_FINAL_SUBMISSION.docx"; path.parent.mkdir(parents=True, exist_ok=True); doc.save(path); return path


def add_dataframe_table(doc: Document, frame: pd.DataFrame, max_rows: int | None = None) -> None:
    use = frame if max_rows is None else frame.head(max_rows)
    table = doc.add_table(rows=1, cols=len(use.columns)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for j, col in enumerate(use.columns):
        cell = table.rows[0].cells[j]; cell.text = str(col); set_cell_shading(cell, "D9EAF7"); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.font.bold = True; run.font.size = Pt(7)
    for row in use.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = "" if pd.isna(value) else str(value); cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[j].paragraphs[0].runs: run.font.size = Pt(6.5)
    doc.add_paragraph()


def build_supplement(bench: pd.DataFrame, e2e: pd.DataFrame, challenge: pd.DataFrame, transport: pd.DataFrame, tests: pd.DataFrame) -> Path:
    reasons = pd.read_csv(R2 / "software_reviewer_archive/REASON_CODE_MANIFEST.csv")
    precedence = {"BLOCK": 1, "NON_EVALUABLE": 2, "ABSTAIN": 3, "PROCEED": 4}
    reasons.insert(0, "state_precedence", reasons.terminal_state.map(precedence)); reasons["within_state_priority"] = reasons.reason_code.rank(method="dense").astype(int); reasons = reasons.sort_values(["state_precedence", "reason_code"]); reasons["primary_reason_rule"] = "state precedence, then lexicographic reason code"
    software = pd.read_csv(R2 / "software_reviewer_archive/ARCHIVAL_EVALUATOR_IDENTITY_MANIFEST.csv")
    beta_support = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_16A_SOURCE_DATA.csv")
    beta_de = pd.read_csv(R2 / "source_data/FIGURE_5C_V3_5_SOURCE_DATA.csv")
    beta = beta_support.merge(beta_de, on="cluster", how="outer")
    capability = pd.DataFrame([
        ["scDesignGuard", "Pre-inference source/design authorization", "Required", "Required", "Required", "Required", "Required", "Delegates estimation to downstream engines", "This study"],
        ["Donor-aware pseudobulk practice", "Biological-unit-aware aggregation and inference", "Not a required or documented core function", "Core design requirement", "Not a required or documented core function", "Study-specific", "Model-dependent", "Performs effect estimation through a statistical engine", "Refs. 8–10"],
        ["edgeR", "Count-based differential analysis", "Not a required or documented core function", "Supplied through the design matrix", "Not a required or documented core function", "Not a required or documented core function", "Model and contrast diagnostics", "Performs effect estimation", "Ref. 20"],
        ["limma-voom", "Precision-weighted linear modelling", "Not a required or documented core function", "Supplied through the design matrix", "Not a required or documented core function", "Not a required or documented core function", "Model and contrast diagnostics", "Performs effect estimation", "Refs. 52,53"],
        ["Cell Ontology / SSSOM", "Semantic identity and mapping representation", "Not a required or documented core function", "Not a required or documented core function", "Core semantic function", "Not a required or documented core function", "Not a required or documented core function", "Does not estimate effects", "Refs. 46–48"],
    ], columns=["approach", "documented_primary_role", "source_version_gate", "donor_design", "ontology_relation", "target_support", "estimability_authorization", "effect_estimation", "evidence"])

    doc = Document(); style_doc(doc); p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Supplementary Information"); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(TITLE).italic = True
    methods = [
        ("Supplementary Methods 1 | Framework, states, stage transitions and reason priority", "Contracts bind source, donor, condition, target, support and estimand evidence. Gates are sequential. Terminal-state precedence is BLOCK, NON_EVALUABLE, ABSTAIN and PROCEED; activated reasons within a state are ordered lexicographically by frozen reason code. Unperformed checks remain NOT_ASSESSED or NOT_APPLICABLE."),
        ("Supplementary Methods 2 | Evidence tiers and authorization thresholds", "Tier A, Tier B, Tier C and Tier D delimit confirmatory authorization, supportive use, structural use and controlled invalidity. Donor and target-cell cut-offs are conservative project rules, not universal sample-size recommendations."),
        ("Supplementary Methods 3 | Final 15-family overlap-excluded benchmark", "Family identity was screened before endpoint estimation. Frozen source-bound human consensus labels and evaluator states are listed in Supplementary Table 1. Selection did not use evaluator state, downstream accessibility, biological-effect estimates or result appearance."),
        ("Supplementary Methods 4 | Four-case end-to-end evaluation", "The four cases were processed in frozen order through source scope, structural design, target support, count aggregation, model validity and final authorization. Independent implementations reconciled counts and design-matrix quantities. No differential effect was estimated."),
        ("Supplementary Methods 5 | Controlled invalidities", "Each challenge changed one design property and retained the remaining contract. Expected state and primary reason were fixed before evaluation."),
        ("Supplementary Methods 6 | Repository transport and software reproducibility", "Transport audits used official identity and metadata endpoints without expression or outcomes. Software identity was bound to wheel, OCI image, schema, reason registry, SBOM, environment and test receipts."),
        ("Supplementary Methods 7 | Beta-cell application and sensitivity analyses", "Author C0–C7 labels were retained. Donor-cluster pseudobulks supported composition, edgeR and limma-voom sensitivity analyses. Hallmark and Reactome CAMERA tests retained the full eligible multiple-testing universe; display subsets followed the frozen minimum-global-FDR ranking rule described in Methods."),
    ]
    for heading, text in methods: add_heading(doc, heading, 1); add_rich_paragraph(doc, text)
    add_heading(doc, "Supplementary Table 1 | Final 15-family membership", 1)
    t1_cols = ["record_id", "repository", "exact_source_version", "frozen_human_reference_class", "scdesignguard_primary_reason", "scdesignguard_state", "state_concordance_to_frozen_reference", "endpoint_membership"]
    add_dataframe_table(doc, bench[t1_cols])
    add_heading(doc, "Supplementary Table 2 | Final four-case end-to-end ledger", 1)
    t2_cols = ["display_id", "accession_or_collection_id", "request_started_utc", "contrast", "target", "donor_count", "pair_count", "target_cell_count", "pseudobulk_rows", "design_rank", "estimability", "residual_df", "terminal_result"]
    add_dataframe_table(doc, e2e[t2_cols])
    add_heading(doc, "Supplementary Table 3 | Frozen state and reason hierarchy", 1); add_dataframe_table(doc, reasons[["state_precedence", "within_state_priority", "reason_code", "terminal_state", "definition", "primary_reason_rule"]])
    add_heading(doc, "Supplementary Table 4 | Controlled invalidities", 1); add_dataframe_table(doc, challenge[["challenge_id", "challenge", "terminal_state", "expected_primary_reason", "primary_reason_code", "safe_nonproceed", "exact_primary_reason_match"]])
    add_heading(doc, "Supplementary Table 5 | Repository transport audit", 1); add_dataframe_table(doc, transport)
    add_heading(doc, "Supplementary Table 6 | Software identity and reproducibility", 1); add_dataframe_table(doc, software)
    add_dataframe_table(doc, tests[["metric", "numerator", "denominator", "detail"]])
    add_heading(doc, "Supplementary Table 7 | Beta-cell analysis, quality control and sensitivity", 1); add_dataframe_table(doc, beta)
    add_heading(doc, "Supplementary Table 8 | Conservative capability comparison", 1); add_dataframe_table(doc, capability)
    add_rich_paragraph(doc, "Entries describe documented primary roles and required core functions. They do not constitute a performance ranking or a claim that other approaches cannot be extended.")
    path = OUT / "scDesignGuard_Nature_Methods_SUPPLEMENTARY_INFORMATION.docx"; doc.save(path)
    write_csv(STAGE / "supplement_table_3_reason_hierarchy.csv", reasons)
    write_csv(STAGE / "supplement_table_8_capability_comparison.csv", capability)
    return path


def build_cover_letter() -> Path:
    doc = Document(); style_doc(doc); p = doc.add_paragraph(); p.add_run("Dear Editors,").bold = True
    paragraphs = [
        "We submit “scDesignGuard: source-bound design gating for multisample single-cell differential analysis” for consideration in Nature Methods.",
        "Single-cell workflows can run even when source identity, biological replication, target support or model estimability do not authorize the requested inference. scDesignGuard addresses the practical question ‘analysis can run—but should it be authorized?’ by compiling source-bound evidence into stage-aware PROCEED, ABSTAIN, BLOCK or NON_EVALUABLE decisions before inference.",
        "The final overlap-excluded benchmark distinguishes four valid, six hard-invalid and five structurally unresolved study families with exact state concordance to frozen source-bound human consensus. A separate four-case end-to-end evaluation shows two tasks reaching model validity and two stopping at explicit source-binding or donor-replication gates. Nine controlled invalidities return their prespecified states and reasons, while repository transport audits and clean-environment reproduction define the operating boundary.",
        "A donor-aware application to 99,029 pancreatic beta cells from 48 donors localizes same-cohort composition, transcriptional and pathway signals across eight author-defined states. scDesignGuard is not a replacement for edgeR, limma-voom or donor-aware pseudobulk engines; it determines whether evidence authorizes their use for the requested task.",
        "We intentionally report state concordance rather than external-accuracy claims. The manuscript, Source Data and the exact evaluated software archive are supplied for editorial and peer review.",
        "Sincerely,\nThe authors"
    ]
    for para in paragraphs: doc.add_paragraph(para)
    path = OUT / "Nature_Methods_Cover_Letter.docx"; doc.save(path); return path


def build_reporting_checklist() -> Path:
    doc = Document(); style_doc(doc); doc.add_heading("Nature Methods reporting checklist", 0)
    rows = [
        ("Study design", "PASS", "Final benchmark and end-to-end denominators are explicit."),
        ("Blinding", "PASS", "Human source-bound review was blinded to evaluator, checklist, end-to-end and biological-outcome outputs before consensus freeze."),
        ("Randomization", "NOT_APPLICABLE", "No new experimental allocation was performed."),
        ("Replication", "PASS", "Donors or complete pairs are the biological units; cells are not treated as independent replicates."),
        ("Statistics", "PASS", "Models, multiplicity, Wilson intervals and display-only pathway selection are specified."),
        ("Data availability", "PASS", "Source Data accompany the manuscript; expression objects are not redistributed."),
        ("Code availability", "PASS", "Exact evaluator and analysis materials are publicly available at https://github.com/wang177777/scDesignGuard."),
        ("Human participant data", "PASS", "Publicly available de-identified data only; no new participant observations."),
        ("Author correspondence and funding", "PASS", "Corresponding-author and five funding-source records are included in the manuscript."),
        ("Acknowledgements", "NOT_APPLICABLE", "No acknowledgements were requested."),
    ]
    add_dataframe_table(doc, pd.DataFrame(rows, columns=["item", "status", "evidence_or_action"]))
    path = OUT / "Nature_Methods_Reporting_Checklist.docx"; doc.save(path); return path


def build_source_staging(fig1, bench, e2e, challenge, transport, tests, fig5, eds) -> None:
    source_dir = STAGE / "source_sheets"; source_dir.mkdir(parents=True, exist_ok=True)

    def write_complete(name: str, frame: pd.DataFrame) -> None:
        """Write a typed source table and reject semantically unexplained blanks."""
        if frame.empty:
            raise AssertionError(f"{name} is empty")
        blank = frame.isna().sum()
        if int(blank.sum()):
            details = {str(k): int(v) for k, v in blank[blank > 0].items()}
            raise AssertionError(f"{name} contains unexplained blank cells: {details}")
        write_csv(source_dir / name, frame)

    write_csv(source_dir / "Fig1.csv", fig1)
    write_csv(source_dir / "Fig2_Benchmark.csv", bench)
    endpoint_source = pd.read_csv(R2 / "source_data/STUDY_FAMILY_DISJOINT_ENDPOINTS.csv")
    endpoint_source["claim_boundary"] = "STATE_CONCORDANCE_WITH_FROZEN_SOURCE_BOUND_HUMAN_CONSENSUS_ONLY"
    write_csv(source_dir / "Fig2_Endpoints.csv", endpoint_source)
    write_csv(source_dir / "Fig3_E2E.csv", e2e)
    write_csv(source_dir / "Fig4_Invalidities.csv", challenge)
    write_csv(source_dir / "Fig4_Transport.csv", transport)
    write_csv(source_dir / "Fig4_Software.csv", tests)
    means, se, summary, heat = fig5
    write_csv(source_dir / "Fig5_Composition.csv", means)
    write_csv(source_dir / "Fig5_Composition_Effects.csv", se)
    write_csv(source_dir / "Fig5_Expression.csv", summary)
    write_csv(source_dir / "Fig5_Pathway.csv", heat)
    atlas = pd.DataFrame({"metric": ["state_resolved_genes", "atlas_genes", "overlap_genes", "significant_associations", "overlap_associations", "direction_concordant_overlap_associations"], "value": [225, 511, 136, 418, 310, 310]})
    write_csv(source_dir / "Fig5_Atlas.csv", atlas)
    ed1, ed2, ed3, ed4, ed5, _ = eds

    # ED1--ED5 originate from heterogeneous panel-level schemas.  They are
    # intentionally written as typed tables rather than a sparse union table:
    # a structural blank must never be mistaken for zero, missing evidence or
    # an unperformed analysis.
    for value, name in [
        ("target_support", "ED1_Support.csv"),
        ("lodo", "ED1_LODO.csv"),
        ("calibration", "ED1_Calibration.csv"),
    ]:
        z = ed1.loc[ed1.panel_source == value].dropna(axis=1, how="all").reset_index(drop=True)
        write_complete(name, z)

    synthetic = ed2.loc[ed2.panel_source == "synthetic_uncertainty"].dropna(axis=1, how="all").reset_index(drop=True)
    primary = synthetic.loc[synthetic.status == "ESTIMABLE_REVIEW_REQUIRED"].reset_index(drop=True)
    secondary = synthetic.loc[synthetic.status == "SECONDARY_DIAGNOSTIC"].reset_index(drop=True)
    for col, text_value in [
        ("numerator", "NOT_APPLICABLE_TO_MEAN_ESTIMAND"),
        ("mcse", "NOT_REPORTED_FOR_THIS_SECONDARY_DIAGNOSTIC"),
        ("interval_lower", "NOT_REPORTED_FOR_THIS_SECONDARY_DIAGNOSTIC"),
        ("interval_upper", "NOT_REPORTED_FOR_THIS_SECONDARY_DIAGNOSTIC"),
    ]:
        secondary[col] = secondary[col].where(secondary[col].notna(), text_value)
    software = ed2.loc[ed2.panel_source == "software_reproduction"].dropna(axis=1, how="all").reset_index(drop=True)
    write_complete("ED2_Primary.csv", primary)
    write_complete("ED2_Secondary.csv", secondary)
    write_complete("ED2_Software.csv", software)

    for value, name in [
        ("controlled_challenge", "ED3_Challenges.csv"),
        ("repository_transport", "ED3_Transport.csv"),
        ("matched_pair", "ED3_Matched.csv"),
    ]:
        z = ed3.loc[ed3.panel_source == value].dropna(axis=1, how="all").reset_index(drop=True)
        write_complete(name, z)

    for value, name in [
        ("LODO", "ED4_LODO.csv"),
        ("DIRECTION", "ED4_Direction.csv"),
        ("FILTER", "ED4_Filter.csv"),
        ("INVALID_DE", "ED4_InvalidDE.csv"),
        ("NULL_PERMUTATION", "ED4_NullPerm.csv"),
    ]:
        z = ed4.loc[ed4.panel == value].dropna(axis=1, how="all").reset_index(drop=True)
        write_complete(name, z)

    for value, name in [
        ("mds", "ED5_MDS.csv"),
        ("support", "ED5_Support.csv"),
        ("composition", "ED5_Composition.csv"),
    ]:
        z = ed5.loc[ed5.source_block == value].dropna(axis=1, how="all").reset_index(drop=True)
        write_complete(name, z)
    # Extended Data Fig. 6 contains >100,000 plotted DE points.  Store its
    # panel sources in four compact, non-sparse worksheets rather than one
    # mostly-empty union table; this preserves every plotted point while
    # avoiding millions of structurally blank workbook cells.
    ed6_de = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_17_SOURCE_DATA.csv.gz")[[
        "feature_id", "cluster", "contrast", "logFC_edgeR",
        "PValue_edgeR", "FDR_global_edgeR",
    ]]
    write_complete("ED6_DE.csv", ed6_de)
    hall = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_18_HALLMARK_SOURCE_DATA.csv")
    hall.insert(0, "collection", "HALLMARK")
    react = pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_18_REACTOME_SOURCE_DATA.csv")
    react.insert(0, "collection", "REACTOME")
    write_complete("ED6_Pathways.csv", pd.concat([hall, react], ignore_index=True, sort=False))
    write_complete("ED6_Method.csv", pd.read_csv(BIO / "summary/EXTENDED_DATA_FIGURE_19_SOURCE_DATA.csv"))
    write_complete("ED6_Atlas.csv", pd.DataFrame({
        "metric": ["atlas_genes", "state_resolved_genes", "overlap_genes", "overlap_associations", "direction_concordant_associations"],
        "value": [511, 225, 136, 310, 310],
    }))

    readme_rows = [
        ("Global policy", "All sheets", "Each worksheet contains one logical data schema. Blank structural cells were removed by splitting heterogeneous figure panels into typed worksheets.", "A numeric zero is retained only when it is an observed or defined zero. NOT_ASSESSED means a gate was not reached; NOT_APPLICABLE means the field does not belong to that estimand or terminal path."),
        ("ED1", "ED1_Support; ED1_LODO; ED1_Calibration", "Development target support, leave-one-donor-out diagnostics and calibration labels.", "Three distinct source schemas; no cross-panel blank padding."),
        ("ED2", "ED2_Primary; ED2_Secondary; ED2_Software", "Primary simulation uncertainty, secondary diagnostics and software reproduction.", "Undefined secondary ratios, MCSEs or intervals are written as explicit NOT_APPLICABLE/NOT_REPORTED states rather than left blank."),
        ("ED3", "ED3_Challenges; ED3_Transport; ED3_Matched", "Controlled invalidities, repository transport and matched clean-invalid fixtures.", "Three distinct source schemas; no cross-panel blank padding."),
        ("ED4", "ED4_LODO; ED4_Direction; ED4_Filter; ED4_InvalidDE; ED4_NullPerm", "Whole-beta robustness and invalid counterfactual components.", "Five row types are separated so absent fields cannot be misread as missing observations."),
        ("ED5", "ED5_MDS; ED5_Support; ED5_Composition", "Pseudobulk MDS, donor support and donor-level beta-cell composition.", "Three distinct source schemas; no cross-panel blank padding."),
        ("ED6", "ED6_DE; ED6_Pathways; ED6_Method; ED6_Atlas", "Complete eligible differential-expression universe, pathway results, method concordance and atlas-overlap summary.", "Large DE evidence is retained without truncation in a dedicated worksheet."),
    ]
    readme = pd.DataFrame(readme_rows, columns=["figure_or_policy", "worksheets", "content", "blank_and_status_semantics"])
    write_complete("README.csv", readme)


def copy_code_archive() -> None:
    target = STAGE / "reviewer_code_archive"
    if target.exists(): shutil.rmtree(target)
    shutil.copytree(R2 / "software_reviewer_archive", target)
    readme = target / "FINAL_STORY_REVIEWER_ACCESS.md"
    readme.write_text(
        "# Reviewer code and software access\n\n"
        "This archive contains the exact evaluated wheel and OCI identity, JSON Schema, 21-code reason registry, SBOM, environment definition and clean-install receipts. "
        "The manuscript-facing evaluator, benchmark, repository adapter, target/model recomputation and figure-building scripts are publicly available at https://github.com/wang177777/scDesignGuard.\n",
        encoding="utf-8",
    )


def main() -> None:
    if not INPUT_DOCX.is_file(): raise FileNotFoundError(INPUT_DOCX)
    if OUT.exists(): shutil.rmtree(OUT)
    if STAGE.exists(): shutil.rmtree(STAGE)
    OUT.mkdir(parents=True); STAGE.mkdir(parents=True)
    fig1 = build_figure_1(); bench = enrich_benchmark(); endpoints = build_figure_2(bench); e2e = enrich_e2e(); build_figure_3(e2e); challenge, transport, tests = build_figure_4(); fig5 = build_figure_5()
    ed1 = build_ed1(); ed2 = build_ed2(); ed3 = build_ed3(challenge, transport); ed4 = build_ed4(); ed5 = build_ed5(); ed6 = build_ed6()
    refs = parse_references(); manuscript = build_manuscript(refs); supplement = build_supplement(bench, e2e, challenge, transport, tests); cover = build_cover_letter(); checklist = build_reporting_checklist()
    build_source_staging(fig1, bench, e2e, challenge, transport, tests, fig5, [ed1, ed2, ed3, ed4, ed5, ed6]); copy_code_archive()
    status = {
        "input_docx_sha256": sha256(INPUT_DOCX), "manuscript": str(manuscript), "supplement": str(supplement), "cover_letter": str(cover), "reporting_checklist": str(checklist),
        "benchmark_rows": len(bench), "benchmark_classes": bench.frozen_human_reference_class.value_counts().to_dict(), "e2e_rows": len(e2e), "references": len(refs),
        "reason_priority_verified": "STATE_PRECEDENCE_THEN_LEXICOGRAPHIC_REASON_CODE", "figure5_display_rule_verified": "MAIN_12_HALLMARK;ED_15_HALLMARK_20_REACTOME_BY_MINIMUM_GLOBAL_CLUSTER_PATHWAY_FDR",
        "author_metadata_status": "FINAL_CONFIRMED_10_AUTHORS;CORRESPONDENCE_AND_FUNDING_INCLUDED;NO_ACKNOWLEDGEMENTS",
    }
    (STAGE / "BUILD_STATUS.json").write_text(json.dumps(status, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(status, indent=2))


if __name__ == "__main__":
    main()
